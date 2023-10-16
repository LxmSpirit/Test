import xlrd
import os
import re
import jieba
import sys
import math
from docx import Document

URL = 'D:/Work/Python/PycharmProjects/untitled/CIMS/ALS/321.xls'
URLfolder = 'D:/Work/Python/PycharmProjects/untitled/CIMS/ALS'
# D:\Work\Python\PycharmProjects\untitled\CIMS\ALS
formaddress = 'D:/Work/Python/PycharmProjects/untitled/Rave-ALS/formdata6.txt'
fieldaddress = 'D:/Work/Python/PycharmProjects/untitled/Rave-ALS/fieldata914.txt'
URLcrf = 'D:/Work/Python/PycharmProjects/untitled/CIMS/CRFdraft/KMHH-PIBC-101_UniqueCRF_V1.0_20230116.docx'
URLcrfsave = 'D:/Work/Python/PycharmProjects/untitled/CIMS/CRFdraft/CRF1-2.docx'
URLcrfsave2 = 'D:/Work/Python/PycharmProjects/untitled/CIMS/CRFfinal/CRF1-919.docx'
URLcrfsave3 = 'D:/Work/Python/PycharmProjects/untitled/CIMS/CRFfinal/CRF1-919-field.docx'
URLlist = []


def readfolder():
    # 遍历指定文件夹下的所有Excel文件并执行分析操作
    for file_name in os.listdir(URLfolder):
        if file_name.endswith('.xls'):
            # 构造输入和输出文件的路径
            # input_path = os.path.join(URLfolder, file_name)
            input_path = URLfolder + "/" + file_name
            URLlist.append(input_path)
    # print(type(input_path))
    # print(input_path)
    # print(URLlist)

    return URLlist


def newreadform(urllist):
    dicform = {}
    for url in urllist:
        workbook = xlrd.open_workbook(url)
        # print(workbook.sheet_names())
        sheet1 = workbook.sheet_by_name('eCRF表单')
        nrows = sheet1.nrows  # 行数
        ncols = sheet1.ncols  # 列数
        # print(nrows)
        for i in range(1, nrows):
            cellOID = sheet1.cell(i, 0).value
            cellName = sheet1.cell(i, 1).value
            # print(cellOID)
            # print(cellName)
            newOID = re.sub(r'\d+', '', cellOID)
            pattern = re.compile('[!@#$%^&*\(\)（）_+[\]{};:,，./<>?\|`~-]')
            cellName2 = re.sub(pattern, '', cellName)
            newName = jieba.lcut_for_search(cellName2)
            s = ""
            for name in newName:
                s = s + " " + name

            if newOID in dicform:
                oldv = dicform[newOID]
                dicform[newOID] = oldv + " " + s
            else:
                dicform[newOID] = s
                # print(cellOID, newOID, cellName, s)
                # print(newOID, s)
        # print(nrows,ncols)
        # print(dicform)
        if "FT" in dicform:
            dicform.pop("FT")
        if "CPH" in dicform:
            dicform.pop("CPH")

    return dicform


def newreadfield(urllist):
    dicfield = {}
    for url in urllist:
        workbook = xlrd.open_workbook(url)
        # print(workbook.sheet_names())
        sheet1 = workbook.sheet_by_name('变量列表')
        nrows = sheet1.nrows  # 行数
        ncols = sheet1.ncols  # 列数
        # print(nrows, ncols)
        i = 17
        while i < nrows:
            line = sheet1.cell(i, 0).value
            # print(line)
            if "--" in line:
                # print("---formOID---")
                formlist = line.split("--")
                formoid2 = formlist[0]
                formoid = re.sub(r'\d+', '', formoid2)
                # print(formoid)
                i = i + 1
                linfieid = sheet1.cell(i, 0).value
                linfiena = sheet1.cell(i, 1).value
                # print("---details---")
                # print(linfieid)
                # print(linfiena)
                while len(linfieid) > 0:
                    i = i + 1
                    if i >= nrows:
                        break
                    linfieid = sheet1.cell(i, 0).value
                    linfiena = sheet1.cell(i, 1).value
                    if len(linfieid) > 0:
                        pattern = re.compile('[!？@#$%^&*\(\)（）_+[\]{};:,，./<>?\|`~-]')
                        fieldName2 = re.sub(pattern, '', linfiena)  # 处理后的fieldname
                        if (formoid in linfieid):
                            newfieldOID = linfieid.replace(formoid, "")
                        else:
                            newfieldOID = linfieid  # 处理后的fieldoid
                        # print(newfieldOID)
                        # print(fieldName2)
                        newName = jieba.lcut_for_search(fieldName2)

                        s = ""
                        for name in newName:
                            s = s + " " + name

                        if newfieldOID in dicfield:
                            oldv = dicfield[newfieldOID]
                            dicfield[newfieldOID] = oldv + " " + s
                        else:
                            dicfield[newfieldOID] = s
            i = i + 1
    return dicfield


def write(data, address):
    file = open(address, mode='w', encoding='utf-8')
    for k, v in data.items():
        file.write(k + '_' + v + '\n')


def coutdic(data):
    keysum = len(data)
    dicv = {}
    onlyv = {}
    dick = {}
    maximum = {}
    Prior = {}
    pv = {}
    for k in data:
        v = data[k]
        vlist = v.split(" ")
        while "" in vlist:
            vlist.remove("")
        for i in vlist:
            str = k + '-' + i
            if str in dicv:
                num = dicv[str]
                dicv[str] = num + 1
            else:
                dicv[str] = 1
        if k in dick:
            numk = dick[k]
            dick[k] = numk + 1
        else:
            dick[k] = len(vlist)
    # 用于计算v
    for k in data:
        # strv = str(v)
        v = data[k]
        vlist = v.split(" ")
        while "" in vlist:
            vlist.remove("")
        for i in vlist:
            if i in onlyv:
                num = onlyv[i]
                onlyv[i] = num + 1
            else:
                onlyv[i] = 1

    sumv = 0
    for a in onlyv:
        sumv = sumv + onlyv[a]

    for a in onlyv:
        num = onlyv[a] / sumv
        pv[a] = num

    for key in dicv:
        n1 = dicv[key]
        k = key.split("-")
        n2 = dick[k[0]]
        res = n1 / n2
        maximum[key] = res

    sum = 0
    for key in dick:
        n1 = dick[key]
        sum = sum + n1

    for key in dick:
        n1 = dick[key]
        res = n1 / sum
        Prior[key] = res

    #终极理解-拉普拉斯修正
    maximum3={} #pxy
    Prior3={} #py
    for key in dicv:
        nxy = dicv[key]#NXY
        k = key.split("-")
        ny = dick[k[0]]#NY
        nx = onlyv[k[1]]#所有x的数量
        res = (nxy + 1) / (ny + nx)
        maximum3[key] = res

    for key in dick:
        n1 = dick[key]
        res = n1 / sum
        Prior3[key] = res

    return maximum3, Prior3, pv, dick, onlyv,sumv

def predict3(pxy, py, px, str, onlyk, onlyv,sumv):
    pattern = re.compile('[!@#$%^&*\(\)（）_+[\]{};:,./<>?\|`~-]')
    newstr1 = re.sub(pattern, '', str)
    newstr = jieba.lcut_for_search(newstr1)
    res = {}
    respxy = {}
    type={}

    for k in pxy:
        k1 = k.split("-")
        y = k1[0]#类别
        x = k1[1]#特征
        if y in type:
            type[y] = type[y] +"-"+x
        else:
            type[y] =  x

    for prestr in newstr:
        for k in type:#类别含有的特征种类
            types = type[k]
            typelist = types.split("-")
            if prestr in typelist:#如果再改类别中找到特征
                pxyk = k+"-"+prestr #pxy的K
                logxyv = math.log(pxy[pxyk]) #pyv = py[k] #pxv = px[prestr]  #pxyv =  pxy[pxyk]
            else:
                logxyv = math.log((1 / sumv)) #pxyv = 1 / sumv
            if k in respxy:
                respxy[k] = respxy[k] +logxyv
            else :
                respxy[k] = logxyv

    for k in respxy:
        if k in res :
            res[k] = res[k] + respxy[k]
        else:
            res[k] = math.log(py[k]) + respxy[k]

    return res

def endres2(res):
    maxres = float("-inf")
    resk = ""
    for k in res:
        if res[k] > maxres:
            maxres = res[k]
            resk = k
    return resk


def removenull(data):
    for k in data:
        # strv = str(v)
        v = data[k]
        vlist = v.split(" ")
        print(vlist)
        while "" in vlist:
            vlist.remove("")
        print(vlist)
        print("---")

def readcrf():#对目标文档进行处理
    doc = Document(URLcrf)
    list = []
    # 读取文档中的内容
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 1' or paragraph.style.name == 'Heading 2':
            list.append(paragraph.text)#获得标题列表

    tables = doc.tables
    a = 0
    for table in tables:
        table.add_row()
        rownum = len(table.rows)
        print(rownum)
        subnum =1
        rows = rownum - 1
        if a < len(list) and len(table.rows[rows].cells)>0:
            table.rows[rownum - 1].cells[0].text = list[a]
            print(list[a])
        while len(table.rows[rows].cells)==0:
            subnum = subnum + 1
            rows  = rownum -subnum
        table.rows[rows].cells[0].text = list[a]
        a = a + 1

    doc.save(URLcrfsave)


def writecrfpar(pxy, py, px, onlyk, onlyv,sumv):
    doc = Document(URLcrfsave)
    listpar = []
    # 读取文档中的内容
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 1' or paragraph.style.name == 'Heading 2':
            listpar.append(paragraph.text)#获得标题列表
    listnew = []
    for par in listpar:
        res = predict3(pxy, py, px, par, onlyk, onlyv,sumv)
        a = endres2(res)
        newpar = par+"_"+a
        listnew.append(newpar)
    print("----")
    print(listnew)
    print(len(listnew))
    print(len(listpar))
    index = 0
    if len(listnew) == len(listpar):
        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1' or paragraph.style.name == 'Heading 2':
                paragraph.text =  listnew[index] # 获得标题列表
                index = index +1

    doc.save(URLcrfsave2)

def writecrffield(pxy, py, px, onlyk, onlyv,sumv):
    doc = Document(URLcrfsave2)
    tables = doc.tables

    for table in tables:
        num = 0
        while num<len(table.rows):
            row = table.rows[num]
            if (len(row.cells) > 0):
                rowtext = row.cells[0].text
                res = predict3(pxy, py, px, rowtext, onlyk, onlyv, sumv)
                resfin = endres2(res)
                newpar = rowtext + ":" + resfin
                row.cells[0].text = newpar
                num = num + 1
            else:
                num = num + 1


    doc.save(URLcrfsave3)


if __name__ == '__main__':
    readcrf()
    urllist = readfolder()
    dicform = newreadform(urllist)
    pxy, py, px, onlyk, onlyv, sumv = coutdic(dicform)
    writecrfpar(pxy, py, px, onlyk, onlyv,sumv)

    #urllist = readfolder()
    #dicform = newreadform(urllist)
    #pxy, py, px, onlyk, onlyv,sumv = coutdic(dicform)
    #str = "实验室"
    #res = predict3(pxy, py, px, str, onlyk, onlyv,sumv)
    #print(str)
    #endres2(res)

    #form 是用-分割的
    dicfield = newreadfield(urllist)
    pxy1, py1, px1, onlyk1, onlyv1, sumv1 = coutdic(dicfield)
    writecrffield(pxy1, py1, px1, onlyk1, onlyv1, sumv1)
    #str = "请详述"
    #res = predict3(pxy1, py1, px1, str, onlyk1, onlyv1,sumv1)
    #print(str)
    #endres2(res)




